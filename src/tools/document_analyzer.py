"""Document analyzer tool for EUC Assessment Agent team.

This module provides a tool for analyzing documents to extract structured information.
"""

import logging
import re
import json
from typing import Dict, List, Optional, Any, Union
from pathlib import Path

logger = logging.getLogger(__name__)


class DocumentAnalyzer:
    """A tool for analyzing documents and extracting structured information."""

    def __init__(self):
        """Initialize the document analyzer tool."""
        # Optional dependencies that will be imported on demand
        self._pdfminer_available = False
        self._docx_available = False
        self._check_dependencies()

    def _check_dependencies(self) -> None:
        """Check if optional dependencies are available."""
        try:
            import pdfminer
            self._pdfminer_available = True
        except ImportError:
            logger.warning("pdfminer.six not available, PDF extraction will be limited")

        try:
            import docx
            self._docx_available = True
        except ImportError:
            logger.warning("python-docx not available, DOCX extraction will be limited")

    def analyze_document(
        self, file_path: Union[str, Path], 
        extract_type: str = "text", 
        target_sections: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """Analyze a document and extract information.
        
        Args:
            file_path: Path to the document file
            extract_type: Type of extraction to perform ("text", "structure", "tables")
            target_sections: Specific sections to extract (e.g., ["Introduction", "Requirements"])
            
        Returns:
            Dictionary with extracted information
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return {"error": f"File not found: {file_path}"}
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == ".pdf":
            return self._analyze_pdf(file_path, extract_type, target_sections)
        elif file_extension == ".docx":
            return self._analyze_docx(file_path, extract_type, target_sections)
        elif file_extension == ".txt":
            return self._analyze_text(file_path, extract_type, target_sections)
        elif file_extension in (".json", ".jsonl"):
            return self._analyze_json(file_path)
        else:
            return {"error": f"Unsupported file type: {file_extension}"}

    def _analyze_pdf(
        self, file_path: Path, 
        extract_type: str, 
        target_sections: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Analyze a PDF document.
        
        Args:
            file_path: Path to the PDF file
            extract_type: Type of extraction to perform
            target_sections: Specific sections to extract
            
        Returns:
            Dictionary with extracted information
        """
        if not self._pdfminer_available:
            try:
                # Fallback to simpler PDF handling
                with open(file_path, "rb") as f:
                    # Read first few bytes to check if it's actually a PDF
                    header = f.read(5)
                    if header != b"%PDF-":
                        return {"error": "Invalid PDF format"}
                    
                    # Read the file in chunks to extract text (very basic)
                    f.seek(0)
                    content = f.read().decode("utf-8", errors="ignore")
                    
                    # Extract visible text (very crude approach)
                    text_content = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]", "", content)
                    
                    return {
                        "content": text_content,
                        "metadata": {
                            "path": str(file_path),
                            "size": file_path.stat().st_size,
                            "extraction_method": "basic",
                            "warning": "Limited extraction due to missing pdfminer.six library"
                        }
                    }
            except Exception as e:
                logger.error(f"Error analyzing PDF: {e}")
                return {"error": f"Error analyzing PDF: {str(e)}"}
        
        try:
            from pdfminer.high_level import extract_text, extract_pages
            from pdfminer.layout import LTTextContainer, LTChar, LTTextBox
            
            if extract_type == "text":
                text = extract_text(file_path)
                
                result = {
                    "content": text,
                    "metadata": {
                        "path": str(file_path),
                        "size": file_path.stat().st_size,
                        "extraction_method": "pdfminer.six"
                    }
                }
                
                # Extract target sections if specified
                if target_sections and text:
                    sections = {}
                    for section in target_sections:
                        # Simple heuristic: Look for the section name followed by text
                        pattern = fr"{re.escape(section)}[^\n]*\n(.*?)(?=\n\s*[A-Z][A-Za-z\s]+\n|\Z)"
                        matches = re.findall(pattern, text, re.DOTALL)
                        if matches:
                            sections[section] = matches[0].strip()
                    
                    if sections:
                        result["sections"] = sections
                
                return result
            elif extract_type == "structure":
                # Extract document structure
                structure = {"pages": []}
                
                for i, page in enumerate(extract_pages(file_path)):
                    page_text = ""
                    text_boxes = []
                    
                    for element in page:
                        if isinstance(element, LTTextContainer):
                            page_text += element.get_text()
                            text_boxes.append({
                                "x0": element.x0,
                                "y0": element.y0,
                                "x1": element.x1,
                                "y1": element.y1,
                                "text": element.get_text().strip()
                            })
                    
                    structure["pages"].append({
                        "page_number": i + 1,
                        "text": page_text,
                        "text_boxes": text_boxes
                    })
                
                return {
                    "structure": structure,
                    "metadata": {
                        "path": str(file_path),
                        "size": file_path.stat().st_size,
                        "extraction_method": "pdfminer.six",
                        "pages": len(structure["pages"])
                    }
                }
            else:
                return {"error": f"Unsupported extraction type for PDF: {extract_type}"}
                
        except Exception as e:
            logger.error(f"Error analyzing PDF with pdfminer: {e}")
            return {"error": f"Error analyzing PDF: {str(e)}"}

    def _analyze_docx(
        self, file_path: Path, 
        extract_type: str, 
        target_sections: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Analyze a DOCX document.
        
        Args:
            file_path: Path to the DOCX file
            extract_type: Type of extraction to perform
            target_sections: Specific sections to extract
            
        Returns:
            Dictionary with extracted information
        """
        if not self._docx_available:
            return {"error": "python-docx library not available"}
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            if extract_type == "text":
                # Extract full text
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                
                result = {
                    "content": full_text,
                    "metadata": {
                        "path": str(file_path),
                        "size": file_path.stat().st_size,
                        "title": doc.core_properties.title or "Untitled",
                        "author": doc.core_properties.author or "Unknown",
                        "sections": len(doc.sections),
                        "paragraphs": len(doc.paragraphs)
                    }
                }
                
                # Extract target sections if specified
                if target_sections and doc.paragraphs:
                    sections = {}
                    current_section = None
                    section_content = []
                    
                    for p in doc.paragraphs:
                        if p.style.name.startswith('Heading') and p.text.strip():
                            # Save previous section if we found it
                            if current_section in target_sections and section_content:
                                sections[current_section] = "\n".join(section_content)
                                section_content = []
                            
                            # Start new section
                            current_section = p.text.strip()
                            section_content = []
                        elif current_section is not None:
                            section_content.append(p.text)
                    
                    # Save last section if needed
                    if current_section in target_sections and section_content:
                        sections[current_section] = "\n".join(section_content)
                    
                    if sections:
                        result["sections"] = sections
                
                return result
            
            elif extract_type == "structure":
                # Extract document structure
                structure = {
                    "paragraphs": [],
                    "tables": []
                }
                
                # Extract paragraphs with styles
                for p in doc.paragraphs:
                    if p.text.strip():
                        structure["paragraphs"].append({
                            "text": p.text,
                            "style": p.style.name
                        })
                
                # Extract tables
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    structure["tables"].append({
                        "table_number": i + 1,
                        "rows": len(table.rows),
                        "columns": len(table.rows[0].cells) if table.rows else 0,
                        "data": table_data
                    })
                
                return {
                    "structure": structure,
                    "metadata": {
                        "path": str(file_path),
                        "size": file_path.stat().st_size,
                        "title": doc.core_properties.title or "Untitled",
                        "author": doc.core_properties.author or "Unknown",
                        "sections": len(doc.sections),
                        "paragraphs": len(doc.paragraphs),
                        "tables": len(doc.tables)
                    }
                }
            
            elif extract_type == "tables":
                # Extract only tables
                tables = []
                
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    tables.append({
                        "table_number": i + 1,
                        "rows": len(table.rows),
                        "columns": len(table.rows[0].cells) if table.rows else 0,
                        "data": table_data
                    })
                
                return {
                    "tables": tables,
                    "metadata": {
                        "path": str(file_path),
                        "size": file_path.stat().st_size,
                        "tables_count": len(tables)
                    }
                }
            
            else:
                return {"error": f"Unsupported extraction type for DOCX: {extract_type}"}
                
        except Exception as e:
            logger.error(f"Error analyzing DOCX: {e}")
            return {"error": f"Error analyzing DOCX: {str(e)}"}

    def _analyze_text(
        self, file_path: Path, 
        extract_type: str, 
        target_sections: Optional[List[str]]
    ) -> Dict[str, Any]:
        """Analyze a plain text document.
        
        Args:
            file_path: Path to the text file
            extract_type: Type of extraction to perform
            target_sections: Specific sections to extract
            
        Returns:
            Dictionary with extracted information
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            result = {
                "content": content,
                "metadata": {
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "lines": content.count('\n') + 1
                }
            }
            
            # Extract target sections if specified
            if target_sections and content:
                sections = {}
                for section in target_sections:
                    # Simple heuristic: Look for section name followed by text
                    pattern = fr"{re.escape(section)}[^\n]*\n(.*?)(?=\n\s*[A-Z][A-Za-z\s]+\n|\Z)"
                    matches = re.findall(pattern, content, re.DOTALL)
                    if matches:
                        sections[section] = matches[0].strip()
                
                if sections:
                    result["sections"] = sections
            
            return result
        
        except Exception as e:
            logger.error(f"Error analyzing text file: {e}")
            return {"error": f"Error analyzing text file: {str(e)}"}

    def _analyze_json(self, file_path: Path) -> Dict[str, Any]:
        """Analyze a JSON document.
        
        Args:
            file_path: Path to the JSON file
            
        Returns:
            Dictionary with extracted information
        """
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                if file_path.suffix.lower() == ".jsonl":
                    # JSON Lines format (one JSON object per line)
                    data = [json.loads(line) for line in f if line.strip()]
                else:
                    # Regular JSON
                    data = json.load(f)
            
            return {
                "data": data,
                "metadata": {
                    "path": str(file_path),
                    "size": file_path.stat().st_size,
                    "format": "jsonl" if file_path.suffix.lower() == ".jsonl" else "json"
                }
            }
        
        except Exception as e:
            logger.error(f"Error analyzing JSON file: {e}")
            return {"error": f"Error analyzing JSON file: {str(e)}"}

    def extract_key_information(self, content: str, info_types: List[str]) -> Dict[str, Any]:
        """Extract specific types of information from text content.
        
        Args:
            content: Text content to analyze
            info_types: Types of information to extract (e.g., ["email", "phone", "url"])
            
        Returns:
            Dictionary with extracted information
        """
        result = {}
        
        for info_type in info_types:
            if info_type == "email":
                # Updated regex to better handle TLD and avoid trailing punctuation
                emails = re.findall(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', content)
                result["emails"] = list(set(emails))
            
            elif info_type == "phone":
                # Multiple phone formats
                phones = re.findall(r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', content)
                result["phones"] = list(set(phones))
            
            elif info_type == "url":
                # Refined regex for www URLs to avoid trailing punctuation
                urls = re.findall(r'https?://[^\s<>\"]+|www\.[a-zA-Z0-9-]+(?:\.[a-zA-Z0-9-]+)*(?:\/[^\s<>\"]*)?', content)
                result["urls"] = list(set(urls))
            
            elif info_type == "date":
                # Various date formats
                dates = re.findall(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}', content)
                # Also match text dates
                text_dates = re.findall(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}', content)
                result["dates"] = list(set(dates + text_dates))
            
            elif info_type == "currency":
                # Match currency values
                currencies = re.findall(r'[$€£¥]\s?\d+(?:[.,]\d+)*|\d+(?:[.,]\d+)*\s?[$€£¥]', content)
                result["currencies"] = list(set(currencies))
        
        return result

    def summarize_document(self, content: str, max_length: int = 200) -> str:
        """Generate a simple summary of document content.
        
        Args:
            content: Document content to summarize
            max_length: Maximum length of the summary
            
        Returns:
            Document summary
        """
        # Simple extractive summarization
        if not content:
            return ""
        
        # Split into sentences
        sentences = re.split(r'(?<=[.!?])\s+', content)
        
        if len(sentences) <= 3:
            # Document is already short, return as is (truncated if needed)
            return content[:max_length] + ("..." if len(content) > max_length else "")
        
        # Extract beginning, important sentences, and end
        important_sentences = [sentences[0]]  # First sentence
        
        # Find sentences with important keywords or capital words (simple heuristic)
        keywords = ["important", "significant", "critical", "main", "key", "primary", "major"]
        middle_sentences = []
        
        for sentence in sentences[1:-1]:
            # Check for capitalized phrases (potential key concepts)
            if re.search(r'\b[A-Z][A-Za-z]{2,}\b', sentence):
                score = 2
            else:
                score = 0
            
            # Check for keywords
            for keyword in keywords:
                if keyword.lower() in sentence.lower():
                    score += 1
            
            # Store sentence with score
            middle_sentences.append((sentence, score))
        
        # Sort by score and take top sentences
        middle_sentences.sort(key=lambda x: x[1], reverse=True)
        important_sentences.extend([s for s, _ in middle_sentences[:2]])
        
        if sentences[-1] not in important_sentences:
            important_sentences.append(sentences[-1])  # Last sentence
        
        # Combine the sentences and ensure max length
        summary = " ".join(important_sentences)
        if len(summary) > max_length:
            # Truncate and add ellipsis, ensuring total length is <= max_length
            summary = summary[:max_length - 3] # Leave space for ellipsis
            summary += "..."
        
        return summary


# Singleton instance
_analyzer = None


def get_document_analyzer() -> DocumentAnalyzer:
    """Get the singleton document analyzer instance."""
    global _analyzer
    if _analyzer is None:
        _analyzer = DocumentAnalyzer()
    return _analyzer 